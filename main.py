import asyncio
from aiogram import Dispatcher, Bot, html, F
from aiogram.client.default import DefaultBotProperties
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.enums import ParseMode
from aiogram.filters import CommandStart
from aiogram.types import ReplyKeyboardRemove, \
    ReplyKeyboardMarkup, KeyboardButton, \
    InlineKeyboardMarkup, InlineKeyboardButton, Message, \
    FSInputFile    
from decouple import config
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path

   
BASE_DIR = Path(__file__).resolve().parent


async def create(list, syn, bot, chat_id):
    template_path = BASE_DIR / "шаблон.docx"
    doc = Document(str(template_path))
    table = doc.tables[0]

    for paragraph in doc.paragraphs:
        if "Синергия" in paragraph.text:
            paragraph.text = paragraph.text.replace("Синергия", f"Синергия {syn}")

    while (len(list) >= len(table.rows)):
        row = table.add_row()
        
        # Настраиваем границы для каждой ячейки в новой строке
        for cell in row.cells:
            tc = cell._tc  # Прямой доступ к XML-элементу ячейки
            
            # Получаем или создаём свойства ячейки (tcPr)
            tcPr = tc.get_or_add_tcPr()
            
            # Создаём элемент границ (если его нет)
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            # Определяем все типы границ
            for border_name in ['top', 'bottom', 'left', 'right']:
                border = tcBorders.find(qn(f'w:{border_name}'))
                if border is None:
                    border = OxmlElement(f'w:{border_name}')
                    tcBorders.append(border)
                
                # Устанавливаем параметры границы
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')


    for i in range(len(list)):
        name = list[i].split(' - ')[0]
        count = list[i].split(' - ')[1].split(' ')[0]
        unit = list[i].split(' - ')[1].split(' ')[1]

        cells = table.rows[i + 1].cells
        cells[0].text = name
        cells[1].text = unit
        cells[2].text = count

        print(f'name: {name}, count: {count}, unit: {unit}')

    time_stamp = datetime.now().strftime("%d_%m_%Y")
    file_path = BASE_DIR / f"Заявка_материалы_Синергия_{syn}_{time_stamp}.docx"
    doc.save(str(file_path))

    document = FSInputFile(path=str(file_path))

    await bot.send_document(
        chat_id=chat_id,
        document=document
    )

class user():
    def __init__(self, user_telegram_id: int):
        self.user_id = user_telegram_id
        self.user_data = False

class userManager():
    def __init__(self):
        self._users = {}

    def get_user(self, user_telegram_id: int) -> user:
        if user_telegram_id not in self._users:
            self._users[user_telegram_id] = user(user_telegram_id)
            print(user_telegram_id, ' - новый пользователь')
        return self._users[user_telegram_id]

token = config('Token')
admin = int(config('Admin'))
um = userManager()
dp = Dispatcher(storage=MemoryStorage())

'''Клавиатуры'''
def create_keyboard_one(user_telegram_id: int):

    kb_list = [
        [KeyboardButton(text="Создать новую заявку"), KeyboardButton(text="Посмотреть мои заявки")],
    ]

    if user_telegram_id == admin:
        kb_list.append([KeyboardButton(text='Админ-панель')])

    keyboard_one = ReplyKeyboardMarkup(
        keyboard=kb_list,
        resize_keyboard=True,
        one_time_keyboard=True,
        input_field_placeholder="Воспользуйтесь меню снизу"
    )

    return keyboard_one

def create_keyboard_two(variable: int):

    kb_list = [
        [KeyboardButton(text="Вернуться на главную")]
    ]

    if variable == 1:
        kb_list.append([KeyboardButton(text="Пример формата заявки")])


    keyboard_one = ReplyKeyboardMarkup(
        keyboard=kb_list,
        resize_keyboard=True,
        one_time_keyboard=True,
        input_field_placeholder="Заявка..."
    )  

    return keyboard_one

def create_keyboard_three(user_telegram_id = None):
    
    kb_list = [
        [KeyboardButton(text="Вернуться на главную")]
    ]

    keyboard_one = ReplyKeyboardMarkup(
        keyboard=kb_list,
        resize_keyboard=True,
        one_time_keyboard=True,
        input_field_placeholder="Выберите необходимое действие"
    )  

    if user_telegram_id != None:
        kb_list.append([KeyboardButton(text='Очистить список моих заявок')])

    return keyboard_one

'''Стартовое сообщение'''
async def send_welcome_message(message: Message):
    await message.answer(f"Привет {message.from_user.full_name}!\nЭтот бот создан для создания заявок на ЦС.", 
                        reply_markup=create_keyboard_one(user_telegram_id=message.from_user.id))   

'''Обработчик стартовой команды'''
@dp.message(CommandStart())
async def handle_text(message: Message) -> None:
    user_telegram_id = message.from_user.id
    um.get_user(user_telegram_id)
    await send_welcome_message(message)

'''Обработчик команды для создания заявки'''
@dp.message(F.text == "Создать новую заявку")
async def handle_text(message: Message):
    await message.answer(f'Отправьте заявку в текстовом формате, придерживаясь формата заявки', 
                        reply_markup=create_keyboard_two(1))


'''Обработчик создания заявки'''
@dp.message(F.text.lower().startswith("заявка"))
async def handle_text(message: Message):

    text = message.text
    lines = text.strip().split('\n')
    syn = lines[1].strip()  # например, "38.1"
    userList2 = [line.strip() for line in lines[2:] if line.strip()]

    await create(userList2, syn, message.bot, message.chat.id)


'''Обработчик команды Пример формата заявки'''
@dp.message(F.text == "Пример формата заявки")
async def handle_text(message: Message):
    await message.answer('''Заявка
38.1

Маркер для CD/DVD Brauberg - 15 шт.
Трубчатый ключ 10/13 230мм - 10 шт.
                         ''', reply_markup=create_keyboard_two(0))

'''Обработчик команды Вернуться на главную'''
@dp.message(F.text == "Вернуться на главную")
async def handle_text(message: Message) -> None:
    await send_welcome_message(message)
    
'''Обработчик команды Посмотреть мои заявки'''
@dp.message(F.text == "Посмотреть мои заявки")
async def handle_text(message: Message):
    if 1+1==2:
        await message.answer('''У вас пока нету заявок''', reply_markup=create_keyboard_three())    
    else:
        await message.answer('''Ваши заявки загружаются. Пожалуйста подождите...''', reply_markup=create_keyboard_three(user_telegram_id=message.from_user.id))

async def main() -> None:
    bot = Bot(token=token, default=DefaultBotProperties(parse_mode=ParseMode.HTML))
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
